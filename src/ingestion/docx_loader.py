import base64
import io
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image

from config import IMAGE_MAX_SIZE
from logger import get_logger
from .base import DocumentChunk

log = get_logger("ingestion.docx")


def load_docx(path: str) -> list[DocumentChunk]:
    chunks: list[DocumentChunk] = []
    source = Path(path).name
    log.info("Opening DOCX: %s", source)

    doc = Document(path)

    # Extract paragraphs
    full_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    log.info("Extracted %d non-empty paragraph(s)", len(full_text))
    if full_text:
        chunks.append(DocumentChunk(
            content="\n".join(full_text),
            chunk_type="text",
            source=source,
            page=1,
            metadata={"file_type": "docx"},
        ))

    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        table_text = "\n".join(rows)
        if table_text.strip():
            chunks.append(DocumentChunk(
                content=table_text,
                chunk_type="table",
                source=source,
                page=1,
                metadata={"file_type": "docx", "table_index": table_idx},
            ))
            log.info("Table %d: extracted %d row(s)", table_idx, len(rows))

    log.info("Extracted %d table(s)", len(doc.tables))

    # Extract inline images
    img_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_bytes = rel.target_part.blob
                img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                img.thumbnail(IMAGE_MAX_SIZE)
                buf = io.BytesIO()
                img.save(buf, format="JPEG", quality=85)
                b64 = base64.b64encode(buf.getvalue()).decode()
                chunks.append(DocumentChunk(
                    content="",
                    chunk_type="image",
                    source=source,
                    page=1,
                    metadata={"file_type": "docx"},
                    image_b64=b64,
                ))
                img_count += 1
                log.debug("Extracted inline image %d", img_count)
            except Exception as e:
                log.warning("Failed to extract image from DOCX rel: %s", e)

    log.info("DOCX load complete | text=%d  tables=%d  images=%d  total_chunks=%d",
             1 if full_text else 0, len(doc.tables), img_count, len(chunks))
    return chunks
